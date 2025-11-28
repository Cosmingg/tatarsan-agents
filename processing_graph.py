from __future__ import annotations

import json
from dataclasses import dataclass, asdict
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, TypedDict

from dotenv import load_dotenv
from langgraph.graph import StateGraph, END

from pydantic import BaseModel, Field
from langchain_core.messages import SystemMessage, HumanMessage
from yandex_gpt_client import ChatYandexGPT


# ---------- 0. Загрузка ТУ ----------

TU_DIR = Path(__file__).with_name("tu")


def load_all_tu_configs() -> Dict[str, Dict[str, Any]]:
    """
    Загружаем все *.json из папки tu/ и строим словарь:
    { tu_id: {"meta": {...}, "data": {...}} }
    где tu_id = значение поля "id" внутри JSON
    """
    configs: Dict[str, Dict[str, Any]] = {}

    if not TU_DIR.exists():
        print(f"[WARN] Папка с ТУ {TU_DIR} не найдена")
        return configs

    for p in TU_DIR.glob("*.json"):
        try:
            with p.open("r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            print(f"[WARN] Ошибка чтения ТУ {p}: {e}")
            continue

        tu_id = str(data.get("id") or p.stem)
        meta_name = f"ТУ {tu_id}"

        configs[tu_id] = {
            "meta": {
                "id": tu_id,
                "name": meta_name,
                "file": str(p),
            },
            "data": data,
        }

    return configs


ALL_TU_CONFIGS = load_all_tu_configs()

DEFAULT_TU_ID = os.getenv("TU_ID")
if DEFAULT_TU_ID not in ALL_TU_CONFIGS and ALL_TU_CONFIGS:
    DEFAULT_TU_ID = next(iter(ALL_TU_CONFIGS))  # первый попавшийся


# ---------- 1. Общее состояние пайплайна ----------


class AppState(TypedDict, total=False):
    # Вход
    file_path: str
    file_ext: str
    file_bytes: bytes

    # Выбранное ТУ
    tu_id: str  # например "3667-013-05608841-2020"

    # Текст документа
    raw_text: str

    # Структурированная заявка
    request_fields: Dict[str, Any]      # как вернул LLM
    matched_items: List[Dict[str, Any]] # найденные позиции в каталоге
    export_payload: Dict[str, Any]

    # Лог шагов для UI
    messages: List[str]


# ---------- 2. Модель заявки (НЭМС) ----------


class RequestFieldsModel(BaseModel):
    """Параметры НЭМС из ТУ, в терминах обозначения."""

    dn_mm: Optional[int] = Field(
        None,
        description="Наружный диаметр патрубков Дн, мм (в обозначении: второе поле, например 325)",
    )
    pressure_kgf_cm2: Optional[float] = Field(
        None,
        description="Рабочее давление, кгс/см² (третье поле обозначения, например 40)",
    )
    length_mm: Optional[int] = Field(
        None,
        description="Длина изделия, мм (четвёртое поле обозначения, например 800)",
    )
    medium_code: Optional[str] = Field(
        None,
        description="Код среды, например 'ВД' — техническая или питьевая вода (пятое поле обозначения)",
    )
    placement_code: Optional[str] = Field(
        None,
        description="Код места размещения на трубопроводе (первая цифра в группе '1-2' и т.п.)",
    )
    connection_code: Optional[str] = Field(
        None,
        description="Код типа соединения с трубопроводом (вторая цифра в группе '1-2', например сварка с наконечником)",
    )
    inner_coating_code: Optional[str] = Field(
        None,
        description="Код внутреннего защитного покрытия (первая цифра в группе '4-3')",
    )
    outer_coating_code: Optional[str] = Field(
        None,
        description="Код наружного защитного покрытия (вторая цифра в группе '4-3')",
    )
    terminals_code: Optional[str] = Field(
        None,
        description="Признак установки клемм (например, 'К' — клеммы установлены, пусто — без клемм)",
    )
    climate_code: Optional[str] = Field(
        None,
        description="Климатическое исполнение по ГОСТ 15150 (например, 'У1', 'УД')",
    )
    notes: Optional[str] = Field(
        None,
        description="Дополнительные требования / комментарии из заявки",
    )


@dataclass
class RequestFields:
    dn_mm: Optional[int] = None
    pressure_kgf_cm2: Optional[float] = None
    length_mm: Optional[int] = None
    medium_code: Optional[str] = None
    placement_code: Optional[str] = None
    connection_code: Optional[str] = None
    inner_coating_code: Optional[str] = None
    outer_coating_code: Optional[str] = None
    terminals_code: Optional[str] = None
    climate_code: Optional[str] = None
    notes: Optional[str] = None


# ---------- 3. LLM (LangChain) ----------

load_dotenv()
yandex_api_key = os.getenv("YANDEX_API_KEY") or os.getenv("YC_API_KEY")
yandex_folder_id = os.getenv("YANDEX_FOLDER_ID") or os.getenv("YC_FOLDER_ID")
if not yandex_api_key or not yandex_folder_id:
    raise RuntimeError(
        "Не найдены переменные YANDEX_API_KEY (или YC_API_KEY) и YANDEX_FOLDER_ID "
        "(или YC_FOLDER_ID). Укажи их в .env перед запуском."
    )

yandex_model = os.getenv("YANDEX_GPT_MODEL", "yandexgpt-lite")
yandex_temperature = float(os.getenv("YANDEX_GPT_TEMPERATURE", "0.1"))
yandex_max_tokens = int(os.getenv("YANDEX_GPT_MAX_TOKENS", "1800"))

llm = ChatYandexGPT(
    model=yandex_model,
    temperature=yandex_temperature,
    max_tokens=yandex_max_tokens,
    api_key=yandex_api_key,
    folder_id=yandex_folder_id,
)

def _extract_text_from_message(message) -> str:
    content = message.content
    if isinstance(content, list):
        parts: List[str] = []
        for block in content:
            if isinstance(block, str):
                parts.append(block)
            elif isinstance(block, dict):
                parts.append(block.get("text") or block.get("json", ""))
            else:
                parts.append(str(block))
        return "\n".join(part for part in parts if part)
    return str(content)


def _parse_request_fields(raw_json: str) -> RequestFieldsModel:
    candidate = raw_json.strip()
    if not candidate:
        raise ValueError("LLM вернул пустой ответ, не могу разобрать RequestFieldsModel.")

    try:
        data = json.loads(candidate)
    except json.JSONDecodeError:
        start = candidate.find("{")
        end = candidate.rfind("}")
        if start == -1 or end == -1:
            raise ValueError(f"Не удалось найти JSON в ответе LLM: {candidate}")
        data = json.loads(candidate[start : end + 1])

    return RequestFieldsModel(**data)


# ---------- 4. Утилиты (OCR/парсинг) ----------


def add_msg(state: AppState, text: str) -> None:
    msgs = state.get("messages") or []
    msgs.append(text)
    state["messages"] = msgs


def read_file_bytes(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def extract_text_from_file(path: str, ext: str, data: bytes) -> str:
    """
    Унифицированное извлечение текста из файла:

    - PNG/JPEG/PDF → Яндекс OCR
    - DOCX → python-docx (параграфы + таблицы)
    - XLS/XLSX → pandas (склеиваем все листы)
    - остальное → пытаемся прочитать как текст
    """
    from docx import Document
    import pandas as pd
    from yandex_ocr_client import recognize_file_to_text, YandexOcrError
    
    p = Path(path)
    ext = ext.lower()

    try:
        # 1) Картинки и PDF — через Яндекс OCR
        if ext in [".png", ".jpg", ".jpeg", ".pdf"]:
            return recognize_file_to_text(str(p))

        # 2) DOCX — читаем текст и таблицы
        if ext == ".docx":
            doc = Document(path)
            parts: List[str] = []

            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = "\t".join(c for c in cells if c)
                    if row_text:
                        parts.append(row_text)

            return "\n".join(parts)

        # 3) Excel — читаем все листы и склеиваем
        if ext in [".xls", ".xlsx"]:
            sheets = pd.read_excel(path, sheet_name=None)
            blocks: List[str] = []
            for sheet_name, df in sheets.items():
                blocks.append(f"### {sheet_name}")
                blocks.append(df.to_string(index=False))
            return "\n\n".join(blocks)

        # 4) Старый .doc — лучше конвертить в PDF/Docx до загрузки
        if ext == ".doc":
            raise YandexOcrError(
                ".doc сейчас не поддерживается. "
                "Попроси клиента прислать PDF или DOCX."
            )

        # 5) Фолбэк — пробуем декоднуть байты как текст
        return data.decode("utf-8", errors="ignore")

    except Exception as e:
        print(f"[extract_text_from_file error] {e}")
        return ""


def match_with_catalog(fields: RequestFields) -> List[Dict[str, Any]]:
    """
    Заглушка сопоставления для НЭМС.
    Сейчас просто собираем читаемое описание из полей и возвращаем один вариант.
    """
    desc_parts: List[str] = []

    if fields.dn_mm is not None:
        desc_parts.append(f"Дн {fields.dn_mm} мм")
    if fields.pressure_kgf_cm2 is not None:
        desc_parts.append(f"PN {fields.pressure_kgf_cm2} кгс/см²")
    if fields.length_mm is not None:
        desc_parts.append(f"L={fields.length_mm} мм")
    if fields.medium_code:
        desc_parts.append(f"среда {fields.medium_code}")
    if fields.placement_code:
        desc_parts.append(f"размещение {fields.placement_code}")
    if fields.connection_code:
        desc_parts.append(f"соединение {fields.connection_code}")
    if fields.inner_coating_code:
        desc_parts.append(f"внутр. покрытие {fields.inner_coating_code}")
    if fields.outer_coating_code:
        desc_parts.append(f"наруж. покрытие {fields.outer_coating_code}")
    if fields.terminals_code:
        desc_parts.append(f"клеммы {fields.terminals_code}")
    if fields.climate_code:
        desc_parts.append(f"климат {fields.climate_code}")

    name = "НЭМС"
    if desc_parts:
        name += " (" + ", ".join(desc_parts) + ")"

    return [
        {
            "item_code": "NEMS-PLACEHOLDER",
            "name": name,
            "score": 0.8,
            "matched_fields": asdict(fields),
        }
    ]


def build_export_payload(fields: RequestFields, matched: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Структура, которую можно отправить в 1С / записать в Excel / CSV.
    """
    return {
        "request_fields": asdict(fields),
        "matched_items": matched,
    }


# ---------- 5. Узлы (агенты) LangGraph ----------


def file_ingestion_node(state: AppState) -> AppState:
    path = state.get("file_path")
    if not path:
        raise ValueError("file_path is missing in state")

    ext = Path(path).suffix.lower()
    data = read_file_bytes(path)

    state["file_ext"] = ext
    state["file_bytes"] = data
    add_msg(state, f"[file_ingestion] Loaded file {Path(path).name} (ext={ext}, size={len(data)} bytes).")
    return state


def text_extraction_node(state: AppState) -> AppState:
    path = state.get("file_path")
    ext = state.get("file_ext")
    data = state.get("file_bytes")

    if not path or ext is None or data is None:
        raise ValueError("file_path/file_ext/file_bytes not set for text extraction")

    needs_ocr = ext in [".png", ".jpg", ".jpeg", ".pdf"]
    if needs_ocr:
        add_msg(
            state,
            f"[ocr] START: отправляю файл {Path(path).name} ({ext}) в Яндекс OCR.",
        )

    text = extract_text_from_file(path, ext, data)
    state["raw_text"] = text
    if needs_ocr:
        status = "OK" if text else "EMPTY"
        add_msg(
            state,
            f"[ocr] RESULT: получено {len(text)} символов (status={status}).",
        )
    add_msg(state, f"[text_extraction] Extracted text of length {len(text)} chars.")
    # 👇 логируем превью распознанного текста (OCR / парсинг)
    preview = text[:500].replace("\n", " ")
    add_msg(state, f"[text_extraction][preview] {preview}")
    return state


def field_extraction_node(state: AppState) -> AppState:
    text = state.get("raw_text", "")
    if not text.strip():
        add_msg(
            state,
            "[field_extraction] raw_text пустой. Сначала загрузите файл или вставьте текст заявки.",
        )
        return state

    # Берём ТУ: либо из состояния (state["tu_id"]), либо дефолтный
    tu_id = state.get("tu_id") or DEFAULT_TU_ID
    tu_cfg = ALL_TU_CONFIGS.get(tu_id)

    if not tu_cfg:
        add_msg(state, f"[field_extraction] TU config '{tu_id}' не найден, работаю без ТУ.")
        tu_json_for_prompt = "{}"
    else:
        tu_json_for_prompt = json.dumps(tu_cfg["data"], ensure_ascii=False, indent=2)
    schema = RequestFieldsModel.model_json_schema()

    system_msg = SystemMessage(
        content=(
            "Ты извлекаешь параметры НЭМС (неразъемное электроизолирующее муфтовое соединение) "
            "из опросного листа/заявки и приводишь их к структуре RequestFieldsModel.\n\n"
            f"Используй в качестве справочника следующие технические условия (ТУ {tu_id}) в формате JSON:\n"
            f"{tu_json_for_prompt}\n\n"
            "Правила:\n"
            "1. Ничего не придумывай — если параметр не указан и не выводится однозначно из ТУ, оставляй null.\n"
            "2. Если давление указано в МПа, можешь подобрать ближайший класс из pressure_classes.\n"
            "3. Если указана среда, сопоставь её с кодом из product_types (МГ, РС, НП, ВД, ТС и т.п.).\n"
            "4. Если указаны явные коды (ВД, У1, цифры покрытий и др.), используй их как есть, сверяясь с JSON ТУ.\n"
            "Возвращай только JSON объект, строго соответствующий схеме RequestFieldsModel."
        )
    )

    user_msg = HumanMessage(
        content=(
            "Вот текст опросного листа/заявки. "
            "Заполни схему RequestFieldsModel, используя JSON с техническими условиями выше.\n"
            "Ответ должен быть только JSON. Если значения отсутствуют, ставь null либо пустую строку.\n\n"
            + text[:6000]
        )
    )

    response_format = {"json_object": {"schema": schema}}
    response = llm.invoke([system_msg, user_msg], response_format=response_format)
    result_json = _extract_text_from_message(response)
    try:
        result_model = _parse_request_fields(result_json)
    except ValueError as exc:
        add_msg(state, f"[field_extraction][error] {exc}")
        raise
    fields = RequestFields(**result_model.dict())
    state["request_fields"] = asdict(fields)
    add_msg(
        state,
        "[field_extraction] Extracted request fields: "
        + json.dumps(asdict(fields), ensure_ascii=False),
    )
    return state


def matching_node(state: AppState) -> AppState:
    if "request_fields" not in state:
        add_msg(state, "[matching] No request_fields in state, nothing to match.")
        return state

    fields = RequestFields(**state["request_fields"])
    items = match_with_catalog(fields)
    state["matched_items"] = items
    add_msg(state, "[matching] Found catalog matches: " + json.dumps(items, ensure_ascii=False))
    return state


def export_node(state: AppState) -> AppState:
    fields_dict = state.get("request_fields", {})
    items = state.get("matched_items", [])

    fields = RequestFields(**fields_dict) if fields_dict else RequestFields()
    payload = build_export_payload(fields, items)
    state["export_payload"] = payload
    add_msg(state, "[export] Built export payload.")
    return state


# ---------- 6. Сборка графа ----------


def build_processing_graph():
    workflow = StateGraph(AppState)

    workflow.add_node("file_ingestion", file_ingestion_node)
    workflow.add_node("text_extraction", text_extraction_node)
    workflow.add_node("field_extraction", field_extraction_node)
    workflow.add_node("matching", matching_node)
    workflow.add_node("export", export_node)

    workflow.set_entry_point("file_ingestion")

    workflow.add_edge("file_ingestion", "text_extraction")
    workflow.add_edge("text_extraction", "field_extraction")
    workflow.add_edge("field_extraction", "matching")
    workflow.add_edge("matching", "export")
    workflow.add_edge("export", END)

    return workflow.compile()


# ---------- 7. Локальный тест ----------

if __name__ == "__main__":
    graph = build_processing_graph()
    example_path = "uploads/example.png"  # подставь свой путь

    init_state: AppState = {
        "file_path": example_path,
        "messages": [],
        # "tu_id": "3667-013-05608841-2020",  # можно указать явно
    }

    final_state = graph.invoke(init_state)
    final_state.pop("file_bytes", None)

    print("=== FINAL STATE ===")
    print(json.dumps(final_state, ensure_ascii=False, indent=2))

    print("\n=== LOG ===")
    for m in final_state.get("messages", []):
        print(m)
